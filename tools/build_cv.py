from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DOCX = Path(r"D:\csme-portfolio\assets\Dr_Victor_Agughasi_CV.docx")

BLUE = RGBColor(31, 78, 121)
BODY = RGBColor(34, 34, 34)
MUTED = RGBColor(80, 80, 80)

PUBLICATIONS = [
    {
        "title": "Firefly-Based Segmentation and Residual Deep Learning for Multi-Class Diabetic Retinopathy Detection",
        "venue": "Inteligencia Artificial",
        "doi": "10.4114/intartif.vol28iss76pp223-252",
    },
    {
        "title": "Enhancing Early Breast Cancer Detection with Infrared Thermography: A Comparative Evaluation of Deep Learning and Machine Learning Models",
        "venue": "Technologies",
        "doi": "10.3390/technologies13010007",
    },
    {
        "title": "Leveraging Transfer Learning for Efficient Diagnosis of COPD Using CXR Images and Explainable AI Techniques",
        "venue": "Inteligencia Artificial",
        "doi": "10.4114/intartif.vol27iss74pp133-151",
    },
    {
        "title": "The Superiority of Fine-tuning over Full-training for the Efficient Diagnosis of COPD from CXR Images",
        "venue": "Inteligencia Artificial",
        "doi": "10.4114/intartif.vol27iss74pp62-79",
    },
    {
        "title": "CX-Net: An Efficient Ensemble Semantic Deep Neural Network for ROI Identification from Chest X-ray Images for COPD Diagnosis",
        "venue": "Machine Learning: Science and Technology",
        "doi": "10.1088/2632-2153/acd2a5",
    },
    {
        "title": "Energy-Efficient Deep Q-Network (EEDQN): Reinforcement Learning-based Approach to Efficient Routing Protocol in Wireless Internet-of-Things",
        "venue": "Indonesian Journal of Electrical Engineering and Computer Science",
        "doi": "10.11591/ijeecs.v33.i2.pp971-980",
    },
    {
        "title": "COPDNet: An Explainable ResNet50 Model for the Diagnosis of COPD from CXR Images",
        "venue": "IEEE INDISCON, 2023",
        "doi": "10.1109/INDISCON58499.2023.10270604",
    },
    {
        "title": "xAI: An Explainable AI Model for the Diagnosis of COPD from CXR Images",
        "venue": "IEEE ICDDS, 2023",
        "doi": "10.1109/ICDDS59137.2023.10434619",
    },
    {
        "title": "i-Net: A Deep CNN Model for White Blood Cancer Segmentation and Classification",
        "venue": "International Journal of Advanced Technology and Engineering Exploration",
        "doi": "10.19101/IJATEE.2021.875564",
    },
    {
        "title": "ResNet-50 vs VGG-19 vs Training from Scratch: A Comparative Analysis of the Segmentation and Classification of Pneumonia from Chest X-ray Images",
        "venue": "Global Transaction Proceedings",
        "doi": "10.1016/j.gltp.2021.08.027",
    },
    {
        "title": "Semi-Supervised Labelling of Chest X-Ray Images Using Unsupervised Clustering for Ground-Truth Generation",
        "venue": "Applied Engineering and Technology",
        "doi": "10.31763/aet.v2i3.1143",
    },
    {
        "title": "Effective Approach for Fine-Tuning Pre-Trained Models for the Extraction of Texts From Source Codes",
        "venue": "ITM Web of Conferences",
        "doi": "10.1051/itmconf/20246503004",
    },
]


def set_font(run, size=10.0, bold=False, italic=False, color=BODY):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_bottom_border(paragraph, color="2F5D8A", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def add_hyperlink(paragraph, text, url, color="0563C1", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)

    if underline:
        underline_el = OxmlElement("w:u")
        underline_el.set(qn("w:val"), "single")
        r_pr.append(underline_el)
    else:
        underline_el = OxmlElement("w:u")
        underline_el.set(qn("w:val"), "none")
        r_pr.append(underline_el)

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_pr.append(r_fonts)

    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), "20")
    r_pr.append(size_el)

    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_page(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)


def base_document():
    doc = Document()
    configure_page(doc.sections[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(10.1)

    if "CV Section" not in doc.styles:
        style = doc.styles.add_style("CV Section", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(13.2)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.0

    return doc


def section_heading(doc, text):
    p = doc.add_paragraph(style="CV Section")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13.2, bold=True, color=BLUE)
    add_bottom_border(p)


def centered_line(doc, text, size=10.4, bold=False, italic=False, color=BODY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)


def body_paragraph(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_font(run, size=10.2, italic=italic)
    return p


def bullet(doc, label, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    label_run = p.add_run(label)
    set_font(label_run, size=10.15, bold=True)
    text_run = p.add_run(text)
    set_font(text_run, size=10.15)
    return p


def number_item(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_font(run, size=10.0)
    return p


def publication_item(doc, index, title, venue, doi):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    intro = p.add_run(f"\"{title}\", {venue}, DOI: ")
    set_font(intro, size=10.0)
    add_hyperlink(p, doi, f"https://doi.org/{doi}")
    return p


def role_entry(doc, title, dates, org, focus="", bullets=None):
    bullets = bullets or []
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.tab_stops.add_tab_stop(Cm(17.7), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(title)
    set_font(r1, size=10.9, bold=True)
    r2 = p.add_run("\t" + dates)
    set_font(r2, size=10.5, italic=True, color=MUTED)

    org_p = doc.add_paragraph()
    org_p.paragraph_format.space_after = Pt(0)
    org_p.paragraph_format.line_spacing = 1.0
    org_run = org_p.add_run(org)
    set_font(org_run, size=10.25)

    if focus:
        focus_p = doc.add_paragraph()
        focus_p.paragraph_format.space_after = Pt(0)
        focus_p.paragraph_format.line_spacing = 1.0
        focus_run = focus_p.add_run(focus)
        set_font(focus_run, size=10.1, italic=True, color=MUTED)

    for item in bullets:
        bullet(doc, "", item)


def build():
    doc = base_document()

    centered_line(doc, "Dr. Agughasi Victor Ikechukwu", size=20, bold=True)
    centered_line(
        doc,
        "Senior Assistant Professor | AI for Healthcare | Medical Imaging | Explainable AI | Multimodal and",
        size=11.2,
        bold=True,
    )
    centered_line(doc, "Agentic AI", size=11.2, bold=True)
    centered_line(
        doc,
        "Department of Computer Science and Medical Engineering, School of Engineering, Dayananda Sagar University, Bangalore, Karnataka-India",
        size=10.5,
    )
    centered_line(
        doc,
        "Mobile: (+91) 7892819690 | Email: victor-csme@dsu.edu.in | Alternate: victor.agughasi@gmail.com",
        size=10.2,
    )
    centered_line(
        doc,
        "Scopus: Author ID 58663521900 | Google Scholar: TQjqJ1UAAAAJ | ORCID: 0000-0002-1175-3089",
        size=10.0,
    )
    centered_line(
        doc,
        "Vidwan ID: 169687 | ResearchGate: Victor-Ikechukwu-Agughasi | LinkedIn: victor-ikechukwu-agughasi",
        size=10.0,
    )

    section_heading(doc, "Professional Summary")
    body_paragraph(
        doc,
        "Senior Assistant Professor in the Department of Computer Science and Medical Engineering, School of Engineering, Dayananda Sagar University, Bangalore, Karnataka-India, with a Ph.D. in Computer Science and experience in artificial intelligence, medical imaging, explainable AI, multimodal machine learning, computer vision, deep learning, clinical decision support, assistive technology, and AI-enabled healthcare systems. Published in Scopus, SCI, and SCIE-indexed journals and IEEE conferences; reviewer for Q1/Q2 journals; mentor of award-winning student innovation teams; and contributor to competitive AI healthcare grants, patents, and institutional AI funding initiatives.",
    )

    section_heading(doc, "Core Competencies and Keywords")
    body_paragraph(
        doc,
        "Artificial Intelligence; Machine Learning; Deep Learning; Computer Vision; Medical Image Analysis; Explainable AI; Multimodal Machine Learning; Agentic AI; Context Engineering; Data Science; Clinical Decision Support; Assistive Technology; Speech and Hearing Technology; Infant Biometrics; Chest X-ray Analysis; COPD Diagnosis; Diabetic Retinopathy; Breast Cancer Thermography; PET-CT Analysis; Radiomics; Python; Java; JavaScript; PHP; MySQL; PostgreSQL; Oracle; Research Methodology; Academic Mentorship; Grant Writing; Patent Development; Sports and Student Leadership.",
    )

    section_heading(doc, "Selected Achievements, Grants, Innovation and Leadership")
    bullet(
        doc,
        "Winner, AIISH Hackathon 2025 on Assistive Technology: ",
        "Project titled \"Development of Interactive Tele-consultation Tool for SLPs and Audiologists\", focused on strengthening remote support for speech-language pathologists and audiologists.",
    )
    bullet(
        doc,
        "Phase 1 Winner / Final-Round Shortlisted Team, AIISH Hack'A'Comm 2026: ",
        "Guided team led by Ms. Afifa Taskeen for \"Adaptive Augmentative and Alternative Communication (AAC) for Children with Cerebral Palsy (CP)\", proceeding to prototype development.",
    )
    bullet(
        doc,
        "VGST K-FIST Level 2 Proposal - INR 30 Lakhs: ",
        "PI for \"Development and Clinical Validation of AI-Assisted Newborn Footprint-Palmprint Biometrics for Infant Identity Assurance\"; shortlisted / accepted for final evaluation on 2 June 2026.",
    )
    bullet(
        doc,
        "AI Centre of Excellence Proposal - INR 20 Crores: ",
        "Institutional proposal shortlisted for final approval to strengthen AI research, innovation, training, and deployment capacity.",
    )
    bullet(
        doc,
        "Three Indian Patent Publications (2026): ",
        "Published patent applications include Victor Ikechukwu Agughasi and Raghavendra K, \"AI-Assisted Newborn Footprint and Palmprint Recognition for Reliable Infant Identity Verification\" (App no. 202641015976, published, waiting for FER); Raghavendra K, Hemanth S R, and Victor Ikechukwu Agughasi, \"Development of an IoT-Enabled Edge-AI System for Real-Time Pest Detection and Non-Lethal Mitigation in Precision Agriculture\" (App no. 202641015258, published, waiting for FER); and Mahesh with Raghavendra K, \"Detection of Congenital Heart Diseases Using Novel Deep Learning Model\" (App no. 202641015562, published).",
    )
    bullet(
        doc,
        "Sports Coordinator and Winner - MITM Mahadasara Sports Utsava 2026: ",
        "Contributed to sports leadership, coordination, student engagement, and institutional team spirit.",
    )

    section_heading(doc, "Education")
    role_entry(
        doc,
        "Ph.D. in Computer Science",
        "Sept. 2019 - July 2024",
        "University of Mysore, Karnataka-India",
        "Thesis: Machine Learning Algorithm for the Diagnosis of Chronic Obstructive Pulmonary Diseases from Chest X-ray Images.",
    )
    role_entry(
        doc,
        "M.Sc. in Computer Science",
        "Apr. 2014 - Mar. 2016",
        "Bangalore University, Bangalore, India",
    )

    section_heading(doc, "Academic and Industry Experience")
    role_entry(
        doc,
        "Senior Assistant Professor, Department of Computer Science and Medical Engineering",
        "June 01, 2026 - Present",
        "School of Engineering, Dayananda Sagar University, Bangalore, Karnataka-India",
        "AI for Healthcare, Medical Image Analysis, Machine Learning, Computer Vision, Explainable AI, Multimodal AI, and Research Mentorship.",
        [
            "Joined DSU to contribute to interdisciplinary teaching, research, innovation, and student mentorship at the interface of computer science, medical engineering, and AI-enabled healthcare.",
            "Focus areas include medical image analysis, explainable AI, infant biometrics, assistive technology, quantum machine learning for healthcare, and clinical decision-support systems.",
        ],
    )
    role_entry(
        doc,
        "Assistant Professor, Department of Computer Science and Engineering (Artificial Intelligence)",
        "Oct. 2021 - May 2026",
        "Maharaja Institute of Technology Mysore, India",
        "Machine Learning, Computer Vision, Big Data Analytics, Digital Image Processing, DBMS, Python for Data Visualization, Research Methodology.",
        [
            "Taught and supervised undergraduate students in AI, machine learning, healthcare AI, computer vision, data analytics, and research-oriented project development.",
            "Guided innovation teams working on assistive technology, AAC systems, EEG analysis, infant biometrics, medical imaging, and applied machine learning systems.",
            "Contributed to institutional research proposals, hackathons, workshops, sports coordination, student mentoring, and academic administration.",
        ],
    )
    role_entry(
        doc,
        "Research Associate",
        "June 2018 - Oct. 2021",
        "Maharaja Institute of Technology Mysore, India",
        "Machine Learning, Big Data Analytics, Machine Learning Projects, Python, Mobile App Development in Java.",
        [
            "Supported teaching and research activities in machine learning, data analytics, Python programming, mobile application development, and student projects.",
            "Developed applied research experience in medical image analysis, explainable AI, and ML-based clinical decision support.",
        ],
    )
    role_entry(
        doc,
        "Visiting Faculty (Voluntary)",
        "Aug. 2018 - May 2019",
        "Dr. Ambedkar Institute for Management Science, Bangalore, India",
        "Information System and Science, Database Management Systems.",
    )
    role_entry(
        doc,
        "Visiting Faculty (Voluntary)",
        "Jul. 2017 - Feb. 2018",
        "St. Aloysius Degree College, Bangalore, India",
        "Information System and Science, Database Management Systems.",
    )
    role_entry(
        doc,
        "Teaching Assistant (Voluntary)",
        "Oct. 2014 - Mar. 2016",
        "St. Joseph's College, Bangalore, India",
        "Computer Fundamentals and Web Design using PHP.",
    )

    section_heading(doc, "Research Interests")
    body_paragraph(
        doc,
        "Medical Imaging; Explainable AI Models; Data Science; Multimodal Machine Learning; Deep Learning; Computer Vision; Context Engineering; Agentic AI Models; Clinical Decision Support; Assistive Technology; AI-Enabled Healthcare Systems; Infant Identity Assurance; Healthcare Data Analytics; Tele-consultation Systems.",
    )

    section_heading(doc, "Intellectual Property and Scholarly Outputs")
    bullet(
        doc,
        "Book manuscript awaiting publication: ",
        "Completed draft of first Quantum ML book, \"Quantum Machine Learning Cookbook for Medical Image Analysis: Recipes, Code Patterns, and Research Pathways for Early-Career Researchers\".",
    )
    bullet(
        doc,
        "Patent publication 1: ",
        "Victor Ikechukwu Agughasi and Raghavendra K, \"AI-Assisted Newborn Footprint and Palmprint Recognition for Reliable Infant Identity Verification\", App no. 202641015976, published (waiting for FER).",
    )
    bullet(
        doc,
        "Patent publication 2: ",
        "Raghavendra K, Hemanth S R, and Victor Ikechukwu Agughasi, \"Development of an IoT-Enabled Edge-AI System for Real-Time Pest Detection and Non-Lethal Mitigation in Precision Agriculture\", App no. 202641015258, published (waiting for FER).",
    )
    bullet(
        doc,
        "Patent publication 3: ",
        "Mahesh and Raghavendra K, \"Detection of Congenital Heart Diseases Using Novel Deep Learning Model\", App no. 202641015562, published.",
    )

    section_heading(doc, "Funded and Shortlisted Research / Innovation Proposals")
    bullet(
        doc,
        "VGST K-FIST Level 2 Proposal - INR 30 Lakhs: ",
        "\"Development and Clinical Validation of AI-Assisted Newborn Footprint-Palmprint Biometrics for Infant Identity Assurance\"; accepted / shortlisted for final evaluation scheduled on 2 June 2026.",
    )
    bullet(
        doc,
        "AI Centre of Excellence Proposal - INR 20 Crores: ",
        "Institutional proposal shortlisted for the final round of approval to establish AI research, innovation, training, and deployment infrastructure.",
    )

    section_heading(doc, "Selected Innovation and Hackathon Projects")
    bullet(
        doc,
        "Development of Interactive Tele-consultation Tool for SLPs and Audiologists: ",
        "Winner, AIISH Hackathon 2025 on Assistive Technology.",
    )
    bullet(
        doc,
        "Adaptive AAC for Children with Cerebral Palsy: ",
        "Phase 1 Winner / Final-Round Shortlisted Team, AIISH Hack'A'Comm 2026; guided team led by Ms. Afifa Taskeen.",
    )
    bullet(
        doc,
        "AI-Assisted Newborn Footprint-Palmprint Biometrics for Infant Identity Assurance: ",
        "PI-led healthcare biometrics research proposal for infant identity assurance.",
    )
    bullet(
        doc,
        "AI-Assisted Multi-Task EEG Analysis for Epilepsy and Related EEG-Based Conditions: ",
        "Student healthcare AI project involving clinical consultation and neurologist guidance.",
    )

    section_heading(doc, "Publications")
    for idx, item in enumerate(PUBLICATIONS, 1):
        publication_item(doc, idx, item["title"], item["venue"], item["doi"])

    section_heading(doc, "Conference Papers Presented")
    bullet(
        doc,
        "",
        "\"Evaluating the Feasibility of a Pre-Processing Framework for Enhanced Text Information Extraction from Source Codes\", ADCIS-2024, BITS Pilani Goa, India.",
    )
    bullet(
        doc,
        "",
        "\"Advances in Thermal Imaging: A Convolutional Neural Network Approach for Improved Breast Cancer Diagnosis\", ICDCOT-2024, SJBIT, India.",
    )
    bullet(
        doc,
        "",
        "\"Optimizing ResNet50 and VGG19 Networks for Accurate COPD Diagnosis Through CXR Analysis\", ERCICAM-2024, NITTE Meenakshi Institute of Technology, India.",
    )
    bullet(
        doc,
        "",
        "\"ResNet-50 vs Training from Scratch: A Comparative Analysis of the Segmentation and Classification of Pneumonia from Chest X-Ray Images\", ICCSA-2021, Acharya Institute of Technology, Bengaluru.",
    )

    section_heading(doc, "Academic Reviewer")
    body_paragraph(
        doc,
        "International Journal of Computer Vision; Computers in Biology and Medicine; Information Sciences; Machine Learning: Science and Technology; Applied Soft Computing Journal; Physics in Medicine and Biology; Indonesian Journal of Electrical Engineering and Computer Science; Digital Health; Inteligencia Artificial; Systems and Soft Computing; Medical Research Archives.",
    )

    section_heading(doc, "Workshops, FDPs, Resource Person Roles and Training")
    bullet(
        doc,
        "2025: ",
        "FDP on \"AI in Medical Imaging and Diagnostics: Current Trends and Challenges\", E&ICT Academy, NIT Patna with IIITDM Jabalpur, IIT Guwahati, MNIT Jaipur, IIT Kanpur, IIT Roorkee, supported by MeitY.",
    )
    bullet(
        doc,
        "2025: ",
        "GIAN FDP on \"Medical Informatics, Radiomics, and Image Analysis for Computer Aided Diagnosis\", NIT Karnataka, Surathkal; presided by Prof. Rangaraj M. Rangayyan, University of Calgary.",
    )
    bullet(
        doc,
        "2024: ",
        "Resource Person, National Level Workshop on Blockchain Technology and Generative AI Models, Ghousia College of Engineering, Ramanagara.",
    )
    bullet(
        doc,
        "2023: ",
        "Resource Person, Two-Day National Workshop on Android Application Development using Java, Maharaja Institute of Technology Mysore.",
    )
    bullet(
        doc,
        "2022: ",
        "Resource Person, Month Internship Program on Deep Learning and its Applications, Hirasugar Institute of Technology; and Three-Day Workshop on Mobile App Development in Android, MIT Thandavapura.",
    )
    bullet(
        doc,
        "2021-2013: ",
        "Resource person / participant in workshops and FDPs on research writing, machine learning, deep learning, data science, data mining, and lateral thinking across India and Nigeria.",
    )

    section_heading(doc, "Merits and Awards")
    bullet(doc, "2026: ", "Sports Coordinator and Winner, MITM Mahadasara Sports Utsava 2026.")
    bullet(doc, "2026: ", "Phase 1 Winner / Final-Round Shortlisted Team, AIISH Hack'A'Comm 2026.")
    bullet(doc, "2025: ", "Winner, AIISH Hackathon on Assistive Technology.")
    bullet(doc, "2024: ", "Best Paper Award, ADCIS-2024, BITS Pilani Goa; Best Paper Award, ERCICAM-2024, NITTE Meenakshi Institute of Technology / IEEE Bengaluru.")
    bullet(doc, "2021: ", "Best Paper Award, ICCSA-2021, Acharya Institute of Technology, Bengaluru.")
    bullet(doc, "2016: ", "Gold Medallist and Best Outgoing Student in PG Science, Computer Science, St. Joseph's College (Autonomous), Bangalore.")
    bullet(doc, "2015-2016: ", "Management Scholarship, St. Joseph's College (Autonomous), Bangalore.")

    section_heading(doc, "Languages, Technical Skills and Interests")
    body_paragraph(
        doc,
        "Languages: English; Kannada (Basic). Programming: Python; Java; JavaScript; PHP. Databases: MySQL; PostgreSQL; Oracle. Technical Skills: Machine Learning; Deep Learning; Computer Vision; Medical Image Analysis; Explainable AI; Data Visualization; Research Methodology; Mobile Application Development; Big Data Analytics; DBMS. Interests: Athletics; Cooking; Cycling; Nature; Swimming; Travelling; Student Mentorship; Community-Oriented AI Innovation; Assistive Technology; Sports Coordination.",
    )

    section_heading(doc, "Academic Referees")
    bullet(
        doc,
        "Dr. Smitha Joyce Pinto: ",
        "Associate Professor, Department of ECE, MIT Mysore, Karnataka - 571477, India; Email: smithapinto_ece@mitmysore.in.",
    )
    bullet(
        doc,
        "Dr. Hemanth S.R.: ",
        "Associate Professor and Head, Department of CSE (Artificial Intelligence), MIT Mysore, Karnataka - 571477, India; Email: hemanthsr_cse@mitmysore.in, hodai@mitmysore.in.",
    )
    bullet(
        doc,
        "Dr. Raghavendra K.: ",
        "Associate Professor, Department of CSE (Artificial Intelligence), MIT Mysore, Karnataka - 571477, India; Email: raghavendrak_ai@mitmysore.in.",
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
